#!/usr/bin/env python2.7

import sys
import os
import requests
import json
import re
import random
from random import randint
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

document=Document()
difficulty=sys.argv[1]

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(20)


nameParagraph = document.add_paragraph('Name: _____________________________')
nameParagraph.style = document.styles['Normal']

for i in range (0,8):
	if difficulty == "easy":
		easyAddition = document.add_paragraph(str(random.randint(0,10)) + ' + ' + str(random.randint(0,10)) + ' = _____ \t\t' + str(random.randint(0,10)) + ' + ' + str(random.randint(0,10)) + ' = ____')

		randA = str(random.randint(0,10))
                randB = str(random.randint(0,10))
		while randB > randA:
	                randB = str(random.randint(0,10))
		randC = str(random.randint(0,10))
                randD = str(random.randint(0,10))
                while randD > randC:
                        randD = str(random.randint(0,10))

		easySubtraction = document.add_paragraph(str(randA) + ' - ' + str(randB) + ' = _____ \t\t' + str(randC) + ' - ' + str(randD) + ' = ____')

        if difficulty == "med":
                easyAddition = document.add_paragraph(str(random.randint(10,40)) + ' + ' + str(random.randint(0,20)) + ' = _____ \t\t' + str(random.randint(10,40)) + ' + ' + str(random.randint(0,20)) + ' = ____')

                randA = str(random.randint(10,40))
                randB = str(random.randint(0,20))
                while randB > randA:
                        randB = str(random.randint(0,20))
                randC = str(random.randint(10,40))
                randD = str(random.randint(0,20))
                while randD > randC:
                        randD = str(random.randint(0,20))

                easySubtraction = document.add_paragraph(str(randA) + ' - ' + str(randB) + ' = _____ \t\t' + str(randC) + ' - ' + str(randD) + ' = ____')

document.add_page_break()


words = ['a', 'about', 'above', 'add', 'after', 'afternoon', 'again', 'air', 'all', 'almost', 'along', 'also', 'always', 'am', 'America', 'an', 'and', 'animal', 'another', 'answer', 'any', 'are', 'around', 'as', 'ask', 'at', 'ate', 'away', 'back', 'be', 'because', 'been', 'before', 'began', 'begin', 'being', 'below', 'between', 'big', 'black', 'book', 'both', 'boy', 'brown', 'but', 'by', 'call', 'came', 'can', 'car', 'carry', 'change', 'children', 'city', 'close', 'come', 'could', 'country', 'cut', 'day', 'did', 'different', 'do', 'does', 'don\'t', 'down', 'each', 'earth', 'eat', 'end', 'enough', 'even', 'every', 'example', 'eye', 'face', 'family', 'far', 'farm', 'father', 'feet', 'few', 'find', 'first', 'fly', 'follow', 'food', 'for', 'found', 'four', 'from', 'get', 'girl', 'give', 'go', 'going', 'good', 'got', 'great', 'group', 'grow', 'had', 'hand', 'hard', 'has', 'have', 'he', 'head', 'hear', 'help', 'her', 'here', 'high', 'him', 'his', 'home', 'house', 'how', 'I', 'idea', 'if', 'important', 'in', 'Indian', 'into', 'is', 'it', 'it\'s', 'just', 'keep', 'kind', 'know', 'land', 'large', 'last', 'late', 'learn', 'leave', 'left', 'let', 'letter', 'life', 'light', 'like', 'line', 'list', 'little', 'live', 'long', 'look', 'made', 'make', 'man', 'many', 'may', 'me', 'mean', 'men', 'might', 'mile', 'miss', 'more', 'most', 'mother', 'mountain', 'move', 'much', 'must', 'my', 'name', 'near', 'need', 'never', 'new', 'next', 'night', 'no', 'not', 'now', 'number', 'of', 'off', 'often', 'oil', 'old', 'on', 'once', 'one', 'only', 'open', 'or', 'other', 'our', 'out', 'over', 'own', 'page', 'paper', 'part', 'people', 'picture', 'place', 'plant', 'play', 'please', 'point', 'pretty', 'put', 'ran', 'read', 'real', 'ride', 'right', 'river', 'round', 'run', 'said', 'same', 'saw', 'say', 'school', 'sea', 'second', 'see', 'seem', 'sentence', 'set', 'she', 'should', 'show', 'side', 'small', 'so', 'some', 'something', 'sometimes', 'song', 'soon', 'sound', 'spell', 'start', 'state', 'still', 'stop', 'story', 'study', 'such', 'take', 'talk', 'tell', 'than', 'thank', 'that', 'the', 'their', 'them', 'then', 'there', 'these', 'they', 'thing', 'think', 'this', 'those', 'thought', 'three', 'through', 'time', 'to', 'together', 'too', 'took', 'tree', 'try', 'turn', 'two', 'under', 'until', 'up', 'us', 'use', 'very', 'walk', 'want', 'was', 'watch', 'water', 'way', 'we', 'well', 'went', 'were', 'what', 'when', 'where', 'which', 'while', 'white', 'who', 'why', 'will', 'with', 'without', 'word', 'work', 'world', 'would', 'write', 'year', 'yes', 'you', 'young', 'your']

table = document.add_table(rows=10, cols=5)

table.style = 'Table Grid'
used = []
print

for i in range (0,10):
	for j in range (0,5):
		randNum = random.randint(0,len(words)-1)
		while used.count(str(words[randNum])) > 0:
			randNum = random.randint(0,len(words)-1)
		used.append(str(words[randNum]))
		cell = table.cell(i,j)
		cell.text = str(words[randNum])
		#cell_font = cell.text_frame.paragraphs[0].runs[0].font
		#cell_font.size = Pt(10)

nameParagraph = document.add_paragraph('\nWrite the following sentences:')

sentences = ['I love my mommy, daddy, and brother.', 'I love my mom because she helps me.', 'I love my mom, she is the best mom ever.', 'My mom is my favorite friend.', 'I like to play baseball.', 'I like to go on adventures with my daddy.', 'I am the best big brother.', 'I will always be a sweetie.', 'I am a Jedi.', 'I am Darth Cenzo.']
used = []

for j in range (0,5):
	randNum = random.randint(0,len(sentences)-1)
	while used.count(str(sentences[randNum])) > 0:
		randNum = random.randint(0,len(sentences)-1)
	used.append(str(sentences[randNum]))
	paragraph = document.add_paragraph(sentences[randNum])
	paragraph = document.add_paragraph('___________________________________')


document.save('test.docx')
